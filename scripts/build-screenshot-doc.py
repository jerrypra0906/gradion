#!/usr/bin/env python3
"""Build DOCX + PDF from UI screenshot manifest."""
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
SHOT_DIR = ROOT / 'docs' / 'ui-screenshots'
MANIFEST = SHOT_DIR / 'manifest.json'
DOCX_OUT = ROOT / 'docs' / 'Gradion_UI_Screenshots.docx'
PDF_OUT = ROOT / 'docs' / 'Gradion_UI_Screenshots.pdf'


def sanitize_pdf_text(text: str) -> str:
    return text.replace('—', '-').replace('’', "'").replace('‘', "'")


def build_docx(manifest):
    doc = Document()
    title = doc.add_heading('GRADION — UI Screenshot Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        'Parent portal walkthrough captured from local environment (localhost:5050). '
        'Test account: parent@gradion.id'
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('Table of Contents', level=1)
    for item in manifest:
        doc.add_paragraph(item['title'], style='List Number')
    doc.add_page_break()

    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'This document provides visual reference screenshots for the Gradion parent-facing '
        'application flows: public landing and onboarding, dashboard navigation, child management, '
        'reports, activity logs, goals, learning modules, and profile settings.'
    )

    for item in manifest:
        doc.add_page_break()
        doc.add_heading(item['title'], level=1)
        p = doc.add_paragraph()
        run = p.add_run('URL: ')
        run.bold = True
        p.add_run(item['url'])

        img_path = SHOT_DIR / item['file']
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.5))
            cap = doc.add_paragraph(f"Figure: {item['title']}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(10)
        else:
            doc.add_paragraph(f'[Missing screenshot: {item["file"]}]')

    doc.save(DOCX_OUT)
    print(f'Wrote {DOCX_OUT}')


def try_pdf(manifest):
    pdf_path = PDF_OUT
    try:
        from fpdf import FPDF
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'fpdf2', '-q'])
        from fpdf import FPDF

    class PDF(FPDF):
        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', align='C')

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 18)
    pdf.cell(0, 12, 'GRADION - UI Screenshot Documentation', ln=True, align='C')
    pdf.set_font('Helvetica', '', 11)
    pdf.cell(0, 8, 'Parent portal walkthrough (localhost:5050)', ln=True, align='C')
    pdf.cell(0, 8, f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
    pdf.ln(8)
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 8, 'Table of Contents', ln=True)
    pdf.set_font('Helvetica', '', 10)
    for item in manifest:
        pdf.cell(0, 6, sanitize_pdf_text(item['title']), ln=True)

    for item in manifest:
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 14)
        pdf.multi_cell(pdf.epw, 8, sanitize_pdf_text(item['title']))
        pdf.set_font('Helvetica', '', 9)
        url_text = item['url'].replace('http://', '').replace('https://', '')
        pdf.multi_cell(pdf.epw, 5, f'URL: {url_text}')
        pdf.ln(3)
        img_path = SHOT_DIR / item['file']
        if img_path.exists():
            # Fit image to page width
            pdf.image(str(img_path), w=190)
        else:
            pdf.cell(0, 8, f"[Missing: {item['file']}]", ln=True)

    pdf.output(str(pdf_path))
    print(f'Wrote {pdf_path}')


def main():
    if not MANIFEST.exists():
        print(f'Missing {MANIFEST}. Run capture-ui-screenshots.mjs first.')
        sys.exit(1)
    manifest = json.loads(MANIFEST.read_text())
    build_docx(manifest)
    try_pdf(manifest)


if __name__ == '__main__':
    main()
