#!/usr/bin/env python3
"""Build DOCX + PDF for child-19 ABA flow screenshots."""
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
SHOT_DIR = ROOT / 'docs' / 'ui-screenshots' / 'child-19'
MANIFEST = SHOT_DIR / 'manifest.json'
DOCX_OUT = ROOT / 'docs' / 'Gradion_Child19_ABA_Flow.docx'
PDF_OUT = ROOT / 'docs' / 'Gradion_Child19_ABA_Flow.pdf'


def sanitize_pdf_text(text: str) -> str:
    return text.replace('—', '-').replace('’', "'").replace('‘', "'")


def build_docx(manifest):
    doc = Document()
    t = doc.add_heading('GRADION — Child Detail & ABA Program Flow', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f'Captured from http://localhost:5050/dashboard/children/19 '
        f'(parent@gradion.id). Generated: {datetime.now():%Y-%m-%d %H:%M}'
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for item in manifest:
        doc.add_heading(item['title'], level=1)
        p = doc.add_paragraph()
        p.add_run('URL: ').bold = True
        p.add_run(item.get('url', ''))
        img = SHOT_DIR / item['file']
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.5))
            cap = doc.add_paragraph(f"Figure: {item['title']}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

    doc.save(DOCX_OUT)
    print(f'Wrote {DOCX_OUT}')


def build_pdf(manifest):
    try:
        from fpdf import FPDF
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'fpdf2', '-q'])
        from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 10, 'GRADION - Child 19 ABA Flow', ln=True, align='C')
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 6, sanitize_pdf_text(datetime.now().strftime('%Y-%m-%d %H:%M')), ln=True, align='C')

    for item in manifest:
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 12)
        pdf.multi_cell(pdf.epw, 7, sanitize_pdf_text(item['title']))
        pdf.set_font('Helvetica', '', 8)
        pdf.multi_cell(pdf.epw, 4, sanitize_pdf_text(item.get('url', '')))
        img = SHOT_DIR / item['file']
        if img.exists():
            pdf.image(str(img), w=190)

    pdf.output(str(PDF_OUT))
    print(f'Wrote {PDF_OUT}')


def main():
    if not MANIFEST.exists():
        print(f'Missing {MANIFEST}')
        sys.exit(1)
    manifest = json.loads(MANIFEST.read_text())
    build_docx(manifest)
    build_pdf(manifest)


if __name__ == '__main__':
    main()
