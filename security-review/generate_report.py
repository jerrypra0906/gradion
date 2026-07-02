import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(r"d:\Claude\Gradion\security-review")
SCREENSHOTS = ROOT / "screenshots"
API_RESULTS = ROOT / "api-test-results.json"
OUT_DOCX = ROOT / "Gradion_Security_Testing_and_Architecture_Review.docx"


def add_h2(doc: Document, text: str):
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str):
    doc.add_heading(text, level=3)


def add_kv(doc: Document, items: list[tuple[str, str]]):
    for k, v in items:
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)


def add_image_if_exists(doc: Document, path: Path, caption: str):
    if not path.exists():
        return
    doc.add_paragraph(caption).runs[0].bold = True
    doc.add_picture(str(path), width=Inches(6.5))


def main():
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc = Document()

    doc.add_heading("Gradion — Security Testing & Architecture Review", level=1)
    add_kv(
        doc,
        [
            ("Generated", now),
            ("Environment", "Local Docker (frontend :5050, backend :5001)"),
            ("Scope", "Public landing + auth + admin CMS + API security controls"),
            ("Method", "OWASP-style manual review + negative/positive API checks + UI evidence screenshots"),
        ],
    )

    add_h2(doc, "1) Executive summary")
    doc.add_paragraph(
        "This report summarizes security testing and an architecture review for Gradion. "
        "It includes evidence-based checks (API responses and UI screenshots) and prioritizes findings "
        "by risk for public release readiness."
    )

    add_h2(doc, "2) Architecture overview (high level)")
    doc.add_paragraph(
        "Gradion is a Next.js (frontend) + Fastify (backend) application backed by Postgres via Prisma. "
        "Authentication uses JWT bearer tokens stored in browser localStorage. "
        "Admin functionality (Users, Landing CMS, CMS editor) is protected via role checks on the backend."
    )
    add_h3(doc, "Security controls observed")
    doc.add_paragraph("- Fastify Helmet enabled with CSP directives (script/frame/img/connect restrictions).")
    doc.add_paragraph("- CORS allowlist with localhost exception.")
    doc.add_paragraph("- Rate limiting enabled globally and for auth endpoints (login/forgot).")
    doc.add_paragraph("- Zod validation used on most request payloads.")
    doc.add_paragraph("- Role-based access control via `requireRole()` middleware.")

    add_h2(doc, "3) UI evidence (screenshots)")
    add_image_if_exists(doc, SCREENSHOTS / "01-landing-desktop.png", "Landing page (desktop)")
    add_image_if_exists(doc, SCREENSHOTS / "02-login.png", "Login page")
    add_image_if_exists(doc, SCREENSHOTS / "03-dashboard.png", "Dashboard after admin login")
    add_image_if_exists(doc, SCREENSHOTS / "04-landing-cms-index.png", "Landing Page CMS index")
    add_image_if_exists(doc, SCREENSHOTS / "05-landing-cms-pricing-edit.png", "Landing Page CMS — Pricing editor")
    add_image_if_exists(doc, SCREENSHOTS / "06-landing-cms-testimonials-edit.png", "Landing Page CMS — Testimonials editor")
    add_image_if_exists(doc, SCREENSHOTS / "07-landing-cms-faq-edit.png", "Landing Page CMS — FAQ editor")

    add_h2(doc, "4) API security test results (positive + negative)")
    if API_RESULTS.exists():
        results = json.loads(API_RESULTS.read_text(encoding="utf-8-sig"))
        for entry in results:
            test = entry.get("test", "")
            res = entry.get("res", {})
            status = str(res.get("status"))
            ok = str(res.get("ok"))
            add_h3(doc, test)
            add_kv(doc, [("Status", status), ("OK", ok)])
            body = res.get("body") or ""
            if len(body) > 1200:
                body = body[:1200] + "…"
            doc.add_paragraph(body)
    else:
        doc.add_paragraph("API test results file was not found.")

    add_h2(doc, "5) Findings and recommendations")
    add_h3(doc, "High priority")
    doc.add_paragraph(
        "- JWT stored in localStorage increases impact of any XSS. Consider httpOnly cookies or strict CSP + zero inline script. "
        "Ensure all user-supplied HTML is sanitized on render."
    )
    doc.add_paragraph(
        "- CORS currently allows any origin containing 'localhost'. In production, remove the localhost wildcard exception."
    )

    add_h3(doc, "Medium priority")
    doc.add_paragraph(
        "- Public CMS pages render sanitized HTML. Keep sanitizer configuration tight and avoid allowing script/iframe unless required."
    )
    doc.add_paragraph(
        "- Rate limiting exists for login; validate rate limiting is configured for all sensitive endpoints (admin create/update, password reset)."
    )

    add_h3(doc, "Low priority / hardening")
    doc.add_paragraph("- Add security headers for frontend delivery (HSTS, referrer policy) at the edge/proxy.")
    doc.add_paragraph("- Add automated dependency scanning (npm audit / Snyk / Dependabot).")
    doc.add_paragraph("- Add audit logging for admin actions (CMS publish, pricing edits, role changes).")

    add_h2(doc, "6) Test plan checklist (recommended ongoing)")
    doc.add_paragraph("Frontend (positive/negative):")
    doc.add_paragraph("- Login/logout flows, incorrect password, locked-out/rate limited scenario.")
    doc.add_paragraph("- Unauthorized navigation to admin pages → should redirect/deny.")
    doc.add_paragraph("- CMS content rendering does not execute scripts (XSS attempt).")
    doc.add_paragraph("Backend API (positive/negative):")
    doc.add_paragraph("- Unauthorized requests to protected endpoints return 401/403.")
    doc.add_paragraph("- IDOR attempts blocked by role/ownership checks.")
    doc.add_paragraph("- Input validation rejects malformed IDs/emails.")

    doc.add_paragraph()
    doc.add_paragraph("End of report.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()

